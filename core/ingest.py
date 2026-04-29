"""
Protocol Document Ingestion for Stromalytix.

Extracts construct parameters from protocol PDFs and DOCX files
using LLM-powered parsing.
"""
import json
import os
import re
from typing import Dict, Optional

from dotenv import load_dotenv

load_dotenv()


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using PyMuPDF (fitz)."""
    import fitz
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    doc.close()
    return text


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx (includes tables)."""
    import io
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract table contents (protocols often have parameter tables)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return "\n".join(paragraphs)


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extract text from plain text files."""
    return file_bytes.decode("utf-8", errors="replace")


def parse_protocol_to_profile(text: str) -> dict:
    """
    Parse protocol text to extract construct profile fields using Claude Haiku.

    Returns dict with keys: target_tissue, cell_types, scaffold_material,
    stiffness_kpa, porosity_percent, cell_density_per_ml,
    experimental_goal, primary_readout.
    None for missing fields.
    """
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        # Fallback: regex-based extraction
        return _regex_extract(text)

    try:
        # Route through rag._build_llm() — prefers LiteLLM/OpenRouter, falls back to Anthropic
        from core.rag import _build_llm
        llm = _build_llm(temperature=0.0, max_tokens=1024)

        prompt = f"""Extract tissue engineering construct parameters from this protocol text.
Return ONLY a JSON object with these exact keys (use null for missing values):

{{
  "target_tissue": "tissue type (e.g. cardiac, liver, neural)",
  "cell_types": ["list", "of", "cell types"],
  "scaffold_material": "material name (e.g. GelMA, collagen)",
  "stiffness_kpa": null or number,
  "porosity_percent": null or number,
  "cell_density_per_ml": null or number,
  "experimental_goal": "disease_modeling | drug_screening | basic_research | null",
  "primary_readout": "viability | contractility | metabolic_activity | gene_expression | null"
}}

Protocol text:
{text[:12000]}

Return ONLY the JSON, no other text."""

        response = llm.invoke(prompt)
        json_str = response.content.strip()
        json_str = re.sub(r"^```json\s*", "", json_str)
        json_str = re.sub(r"\s*```$", "", json_str)
        return json.loads(json_str)

    except Exception as e:
        print(f"LLM extraction failed: {e}, falling back to regex")
        return _regex_extract(text)


def _regex_extract(text: str) -> dict:
    """Fallback regex-based parameter extraction."""
    result = {
        "target_tissue": None,
        "cell_types": None,
        "scaffold_material": None,
        "stiffness_kpa": None,
        "porosity_percent": None,
        "cell_density_per_ml": None,
        "experimental_goal": None,
        "primary_readout": None,
    }

    lower = text.lower()

    # Tissue type
    tissues = ["cardiac", "liver", "hepatic", "neural", "brain", "bone",
               "cartilage", "skin", "lung", "kidney", "tumor", "vascular",
               "muscle", "intestinal", "pancreatic"]
    for tissue in tissues:
        if tissue in lower:
            result["target_tissue"] = tissue
            break

    # Stiffness
    match = re.search(r"(\d+\.?\d*)\s*kPa", text, re.IGNORECASE)
    if match:
        result["stiffness_kpa"] = float(match.group(1))

    # Porosity
    match = re.search(r"(\d+\.?\d*)\s*%\s*(?:porosity|porous)", text, re.IGNORECASE)
    if match:
        result["porosity_percent"] = float(match.group(1))

    # Cell density
    match = re.search(r"(\d+\.?\d*)\s*[×x]\s*10\^?(\d+)\s*(?:cells?/mL|cells?/ml)", text, re.IGNORECASE)
    if match:
        result["cell_density_per_ml"] = float(match.group(1)) * (10 ** int(match.group(2)))

    # Scaffold material
    materials = ["GelMA", "collagen", "alginate", "fibrin", "PEGDA",
                 "hyaluronic acid", "Matrigel", "silk", "PCL", "PLGA"]
    for mat in materials:
        if mat.lower() in lower:
            result["scaffold_material"] = mat
            break

    # Cell types
    cell_types = []
    cell_patterns = ["cardiomyocyte", "fibroblast", "hepatocyte", "neuron",
                     "astrocyte", "endothelial", "epithelial", "stem cell",
                     "iPSC", "MSC", "osteoblast", "chondrocyte"]
    for ct in cell_patterns:
        if ct.lower() in lower:
            cell_types.append(ct)
    if cell_types:
        result["cell_types"] = cell_types

    # Goal
    if "disease model" in lower or "disease" in lower:
        result["experimental_goal"] = "disease_modeling"
    elif "drug screen" in lower:
        result["experimental_goal"] = "drug_screening"

    return result
