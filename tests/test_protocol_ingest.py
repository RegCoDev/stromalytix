"""Tests for protocol document ingestion (core/ingest.py)."""

import json
from unittest.mock import patch, MagicMock

from core.ingest import (
    extract_text_from_pdf,
    extract_text_from_docx,
    extract_text_from_txt,
    parse_protocol_to_profile,
    _regex_extract,
)


# --- Text extraction tests ---

def test_extract_text_from_txt():
    """Plain text files are read directly."""
    content = "GelMA 6% scaffold at 10 kPa for cardiac tissue"
    text = extract_text_from_txt(content.encode("utf-8"))
    assert "GelMA" in text
    assert "cardiac" in text


def test_extract_text_from_pdf():
    """PDF extraction works with PyMuPDF."""
    import fitz
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Hepatic model using GelMA 4% at 3.5 kPa")
    pdf_bytes = doc.tobytes()
    doc.close()

    text = extract_text_from_pdf(pdf_bytes)
    assert "Hepatic" in text
    assert "GelMA" in text


def test_extract_text_from_docx():
    """DOCX extraction works with python-docx."""
    from docx import Document
    import io
    doc = Document()
    doc.add_paragraph("Neural tissue model")
    doc.add_paragraph("Scaffold: collagen I at 0.5 kPa")
    doc.add_paragraph("Cell type: iPSC-derived neurons")
    buf = io.BytesIO()
    doc.save(buf)

    text = extract_text_from_docx(buf.getvalue())
    assert "Neural" in text
    assert "collagen" in text
    assert "iPSC" in text


def test_docx_table_extraction():
    """Tables in DOCX are extracted."""
    from docx import Document
    import io
    doc = Document()
    doc.add_paragraph("Protocol Parameters")
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Material"
    table.cell(0, 1).text = "GelMA 6%"
    table.cell(1, 0).text = "Stiffness"
    table.cell(1, 1).text = "7.5 kPa"
    table.cell(2, 0).text = "Cell Density"
    table.cell(2, 1).text = "5 million/mL"
    buf = io.BytesIO()
    doc.save(buf)

    text = extract_text_from_docx(buf.getvalue())
    assert "GelMA" in text
    assert "7.5 kPa" in text
    assert "5 million" in text


# --- Regex fallback extraction tests ---

def test_regex_extract_tissue():
    """Regex extractor finds tissue type."""
    result = _regex_extract("We cultured cardiac tissue models using GelMA at 10 kPa")
    assert result["target_tissue"] == "cardiac"


def test_regex_extract_stiffness():
    """Regex extractor finds stiffness."""
    result = _regex_extract("The hydrogel stiffness was 7.5 kPa")
    assert result["stiffness_kpa"] == 7.5


def test_regex_extract_material():
    """Regex extractor finds scaffold material."""
    result = _regex_extract("Cells were encapsulated in GelMA hydrogels")
    assert result["scaffold_material"] == "GelMA"


def test_regex_extract_cell_types():
    """Regex extractor finds cell types."""
    result = _regex_extract("Co-culture of cardiomyocyte and fibroblast populations")
    assert "cardiomyocyte" in result["cell_types"]
    assert "fibroblast" in result["cell_types"]


def test_regex_extract_goal():
    """Regex extractor finds experimental goal."""
    result = _regex_extract("for disease modeling applications")
    assert result["experimental_goal"] == "disease_modeling"


# --- LLM extraction tests (mocked) ---

def test_parse_protocol_with_mock_llm():
    """LLM extraction returns dict with parameters."""
    mock_data = {
        "target_tissue": "cardiac",
        "cell_types": ["cardiomyocytes"],
        "scaffold_material": "GelMA",
        "stiffness_kpa": 10.0,
        "porosity_percent": None,
        "cell_density_per_ml": 5000000,
        "experimental_goal": "disease_modeling",
        "primary_readout": "contractility",
    }
    mock_response = MagicMock()
    mock_response.content = json.dumps(mock_data)
    mock_llm = MagicMock()
    mock_llm.invoke.return_value = mock_response

    with patch.dict("os.environ", {"ANTHROPIC_API_KEY": "test-key"}), \
         patch("langchain_anthropic.ChatAnthropic", return_value=mock_llm):
        result = parse_protocol_to_profile("Cardiac protocol using GelMA at 10 kPa")

    assert result["target_tissue"] == "cardiac"
    assert result["scaffold_material"] == "GelMA"
    assert result["stiffness_kpa"] == 10.0


def test_parse_protocol_no_api_key_uses_regex():
    """Falls back to regex when no API key."""
    import os
    old = os.environ.pop("ANTHROPIC_API_KEY", None)
    try:
        result = parse_protocol_to_profile("Hepatic model using collagen at 3.5 kPa")
        assert result["target_tissue"] == "hepatic"
        assert result["scaffold_material"] == "collagen"
        assert result["stiffness_kpa"] == 3.5
    finally:
        if old:
            os.environ["ANTHROPIC_API_KEY"] = old
